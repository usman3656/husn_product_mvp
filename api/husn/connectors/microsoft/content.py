"""Download + extract text from OneDrive files.

Microsoft Graph returns binary content via /items/{id}/content (or
/drives/{drive}/items/{id}/content for shared items). We download bytes,
parse per extension, return plain text. Capped at MAX_DOWNLOAD_BYTES so a
single huge file doesn't blow up the worker.

Extracted text is stashed on the raw_artifact payload as `_extracted_text`
(JSON-safe, single source of truth). The normalizer picks it up as body.
"""

from __future__ import annotations

import csv as csvlib
import io
from html.parser import HTMLParser
from typing import Any

from husn.connectors.microsoft.client import MicrosoftClient
from husn.core.logging import log

MAX_DOWNLOAD_BYTES = 5_000_000  # 5 MB — skip larger files for now
MAX_EXTRACTED_CHARS = 8_000  # ceiling on what lands in artifact.body
MAX_XLSX_ROWS_PER_SHEET = 50
MAX_XLSX_SHEETS = 6


_TEXT_EXTENSIONS = {"html", "htm", "csv", "txt", "md", "log"}
_OFFICE_EXTENSIONS = {"docx", "xlsx"}  # pptx deferred
_ALL_EXTRACTABLE = _TEXT_EXTENSIONS | _OFFICE_EXTENSIONS


def _ext(name: str | None) -> str:
    if not name:
        return ""
    if "." not in name:
        return ""
    return name.rsplit(".", 1)[-1].lower()


def is_extractable(item: dict[str, Any]) -> bool:
    """Cheap test BEFORE downloading: should we bother?"""
    if not item.get("file"):
        return False
    size = item.get("size") or 0
    if size > MAX_DOWNLOAD_BYTES:
        return False
    ext = _ext(item.get("name"))
    return ext in _ALL_EXTRACTABLE


def _content_url(item: dict[str, Any]) -> str:
    """Pick /me/drive or /drives/{drive}/items endpoint based on where the
    file lives. Shared items have a non-self driveId in parentReference.
    """
    file_id = item["id"]
    parent_ref = item.get("parentReference") or {}
    drive_id = parent_ref.get("driveId")
    if drive_id:
        return f"/drives/{drive_id}/items/{file_id}/content"
    return f"/me/drive/items/{file_id}/content"


class _HtmlStripper(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._chunks: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in ("script", "style"):
            self._skip += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in ("script", "style") and self._skip > 0:
            self._skip -= 1

    def handle_data(self, data: str) -> None:
        if self._skip == 0:
            self._chunks.append(data)

    def text(self) -> str:
        return " ".join("".join(self._chunks).split())


def _extract_html(data: bytes) -> str:
    parser = _HtmlStripper()
    parser.feed(data.decode("utf-8", errors="replace"))
    return parser.text()


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8-sig", errors="replace")  # utf-8-sig strips Excel BOM
    out: list[str] = []
    try:
        reader = csvlib.reader(io.StringIO(text))
        for row in reader:
            if not row:
                continue
            line = " | ".join(c.strip() for c in row if c is not None)
            if line:
                out.append(line)
            if len(out) >= 200:
                out.append("…")
                break
    except Exception:
        return text
    return "\n".join(out)


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_docx(data: bytes) -> str | None:
    try:
        from docx import Document
    except ImportError:
        log.warning("husn.microsoft.content.docx.unavailable")
        return None
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        log.info("husn.microsoft.content.docx.parse_failed", err=str(e)[:80])
        return None
    lines: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_xlsx(data: bytes) -> str | None:
    try:
        from openpyxl import load_workbook
    except ImportError:
        log.warning("husn.microsoft.content.xlsx.unavailable")
        return None
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as e:
        log.info("husn.microsoft.content.xlsx.parse_failed", err=str(e)[:80])
        return None
    out: list[str] = []
    for ws in wb.worksheets[:MAX_XLSX_SHEETS]:
        out.append(f"## Sheet: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= MAX_XLSX_ROWS_PER_SHEET:
                out.append("…")
                break
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


async def extract_text(mc: MicrosoftClient, item: dict[str, Any]) -> str | None:
    """Download + parse OneDrive file. Returns plain text or None if we skip / fail.

    Result is capped at MAX_EXTRACTED_CHARS so the artifact.body stays reasonable.
    """
    if not is_extractable(item):
        return None

    ext = _ext(item.get("name"))
    url = _content_url(item)

    try:
        # Use .request() so we get the raw bytes; .get() expects JSON.
        r = await mc.request("GET", url)
    except Exception:
        log.exception("husn.microsoft.content.download_failed", file_id=item.get("id"))
        return None
    if r.status_code == 302:
        # Some endpoints return 302 to a CDN; httpx follows by default but be defensive
        try:
            r = await mc._client.get(r.headers["Location"])
        except Exception:
            return None
    if r.status_code >= 300:
        log.info(
            "husn.microsoft.content.download_status",
            status=r.status_code,
            file_id=item.get("id"),
        )
        return None
    data = r.content
    if len(data) > MAX_DOWNLOAD_BYTES:
        return None

    try:
        if ext in ("html", "htm"):
            text = _extract_html(data)
        elif ext == "csv":
            text = _extract_csv(data)
        elif ext in ("txt", "md", "log"):
            text = _extract_txt(data)
        elif ext == "docx":
            text = _extract_docx(data)
        elif ext == "xlsx":
            text = _extract_xlsx(data)
        else:
            return None
    except Exception:
        log.exception("husn.microsoft.content.extract_failed", file_id=item.get("id"), ext=ext)
        return None

    if not text:
        return None
    text = text.strip()
    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS] + "…"
    return text
